import os
import sys
import shutil
from os import path
from importlib.machinery import SourceFileLoader
import docx
import time
import threading

import docx.enum
import docx.enum.text

FontName = "Arial"
FontSize = 11
BigFontSize = 15

ProjectPath = sys.path[0]
ReviewArchive = path.join(ProjectPath, "ReviewArchive")
NewFolder = path.join(ProjectPath, "CompiledReview")

FileIgnore = [
	"__pycache__",
	"SelfLearning",
]

Rename = {
	"XXthCenturyMusic": "20th Century Music",
	"I_A": "1A",
	"I_B": "1B",
}

Append = [
	
]

SkippedFiles = []
Compiling = 0

def MakeDir(Directory):
	if not path.exists(Directory):
		os.makedirs(Directory)

def DeleteDir(Directory):
	if path.exists(Directory):
		shutil.rmtree(Directory)

def CopyFile(FilePath, NewPath):
	try:
		shutil.copy(FilePath, NewPath)	
	except:
		print(f'{FilePath} failed to parse.')

def FixPascalCase(String):
	NewString = ''.join(' ' + Character.lower() if Character.isupper() else Character for Character in String)

	if NewString.startswith(" "):
		NewString = NewString[1:]

	return NewString.title()

def AddParagraph(Document, Paragraph, ItemDirectory):
	ParagraphLines = 0

	try:
		if "#" in open(ItemDirectory, "r").read():
			print("Item has comments: " + ItemDirectory)

		ReviewNotes = SourceFileLoader("ReviewCompiler", ItemDirectory).load_module()
		Answers = ReviewNotes.Answers

		LastParagraphType = 0

		if len(Answers) == 0: 
			raise ValueError('Module is empty')
			
		for List in Answers:
			ParagraphLines += 1

			if isinstance(Answers, str) and Answers.startswith("poem"):
				PoemParts = Answers.split("\n\n")
				PoemInfo = PoemParts[0].split("\n")

				Paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
				
				Title = Paragraph.add_run(PoemInfo[1] + "\n")
				Title.bold = True
				Title.font.size = docx.shared.Pt(BigFontSize) 

				for ExtraPoemInfo in PoemInfo[2:]:
					Paragraph.add_run(ExtraPoemInfo + "\n").bold = True
				
				Paragraph.add_run("\n")

				for Stanza in PoemParts[1:]:
					Paragraph.add_run(Stanza.lstrip() + "\n\n")

				break

			ParagraphType = (isinstance(List, str) and 1) or (isinstance(List[0], str) and 2) or (isinstance(List[0], list) and 3)

			if LastParagraphType != 0 and LastParagraphType != ParagraphType:
				Paragraph.add_run("\n")
			LastParagraphType = ParagraphType

			if ParagraphType == 1:
				Paragraph.add_run(List + "\n")
			elif ParagraphType == 2:
				Paragraph.add_run(List[0]).bold = True
				Paragraph.add_run(f' - {List[1]}\n').bold
			elif ParagraphType == 3:
				Paragraph.add_run(List[1] + "\n").bold = True
				
				for Answer in List[0]:
					Paragraph.add_run(f'- {Answer}\n')
					ParagraphLines += 1
			else:
				raise ValueError('List is not identified')
			
		return ParagraphLines
	except Exception as Error:
		print("Failed to load module,", Error, ItemDirectory)
		return False

def Scan(Directory):
	global Compiling
	Compiling += 1
	DirectoryItems = os.listdir(Directory)

	for Item in DirectoryItems:
		ItemDir = path.join(Directory, Item)
		ArchiveDir = path.join(NewFolder, path.relpath(ItemDir, ReviewArchive))

		if Item in FileIgnore or ItemDir in SkippedFiles: 
			continue

		if not "." in Item:
			MakeDir(ArchiveDir)
			threading.Thread(target=Scan, args=(ItemDir,)).start()
			continue

		if not Item.endswith(".py"):
			CopyFile(ItemDir, ArchiveDir)
			continue

		Document = docx.Document()
		Style = Document.styles['Normal']

		Font = Style.font
		Font.name = FontName
		Font.size = docx.shared.Pt(FontSize)

		Paragraph = Document.add_paragraph()
		Paragraph.style = Document.styles['Normal']

		NewItemName = None
		Modules = []

		for FilesToAppend in Append:
			if not Item in FilesToAppend[1:]:
				continue

			NewItemName = FilesToAppend[0]

			for TargetFileName in FilesToAppend:
				if TargetFileName == NewItemName:
					continue

				TargetFile = path.join(Directory, TargetFileName)

				if not path.exists(TargetFile): 
					print(f'{TargetFileName} does not exist in {Directory}')

					Modules = []
					NewItemName = None
					break
				
				Modules.append(TargetFile)

		if len(Modules) == 0:
			EnumCopy = ItemDir[:-3] + "Enum.py"
			Modules = [ItemDir]

			if path.exists(EnumCopy):
				Modules.append(EnumCopy)

		for Module in Modules:
			SkippedFiles.append(Module)

			AddParagraph(Document, Paragraph, Module)
			Paragraph.add_run("\n")

		if NewItemName == None:
			NewItemName = Item[:-3]

			if NewItemName in Rename:
				NewItemName = Rename[NewItemName]
			else:
				NewItemName = FixPascalCase(NewItemName)
		
		NewFileDir = path.join(Directory, NewItemName)
		NewFileDir = path.relpath(NewFileDir, ReviewArchive)
		NewFileDir = path.join(NewFolder, NewFileDir)
		NewFileDir = NewFileDir + ".docx"

		Document.save(NewFileDir)
	
	Compiling -= 1

if path.exists(NewFolder):
	DeleteDir(NewFolder)
	print("Deleted compiled folder")
	sys.exit(0)

MakeDir(NewFolder)

Start = time.time()
Scan(ReviewArchive)

while Compiling > 0:
	time.sleep(0)

End = time.time()

print("Finished compiling!")
print(f'Compiling took {str(End - Start)[:4]}s')

UploadToDrive = None

while UploadToDrive == None:
	UploadToDrive = input("Upload to drive? [y/n] ").lower()
	UploadToDrive = (UploadToDrive == "y" and True) or (UploadToDrive == "n" and False) or None

from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

DriveService = build('drive', 'v3', credentials=InstalledAppFlow.from_client_secrets_file(
	{
		"installed": {
			"client_id": ".apps.googleusercontent.com",
			"project_id": "",
			"auth_uri": "https://accounts.google.com/o/oauth2/auth",
			"token_uri": "https://oauth2.googleapis.com/token",
			"auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
			"client_secret": "",
			"redirect_uris": ["http://localhost"]
		}
	}, 
	['https://www.googleapis.com/auth/drive']
).run_local_server(port=0))

def UploadFolder(FolderPath, FolderId):
	DriveFiles = DriveService.files()

	FolderName = os.path.basename(FolderPath)
	NewDriveFolder = DriveFiles.list(
		q=f"name='{FolderName}' and mimeType='application/vnd.google-apps.folder' and parents in '{FolderId}' and trashed = false",
	).execute()['files']

	if len(NewDriveFolder) > 0:
		NewDriveFolder = NewDriveFolder[0]['id']
	else:
		NewDriveFolder = DriveFiles.create(body={
			'name': FolderName,
			'mimeType': 'application/vnd.google-apps.folder',
			'parents': [FolderId]
		}, fields='id').execute().get('id')

	for ItemName in os.listdir(FolderPath):
		ItemPath = os.path.join(FolderPath, ItemName)
		
		if os.path.isdir(ItemPath):
			UploadFolder(ItemPath, NewDriveFolder) 
		elif os.path.isfile(ItemPath):
			ExistingFile = DriveFiles.list(
				q=f"name='{ItemName}' and parents in '{NewDriveFolder}' and trashed = false",
			).execute()['files']

			if len(ExistingFile) > 0:
				DriveFiles.update(
					fileId = ExistingFile[0]['id'],
					media_body=MediaFileUpload(ItemPath, resumable=True)
				).execute()
			else:
				DriveFiles.create(
					body={
						'name': ItemName,
						'parents': [NewDriveFolder]
					}, 
					media_body=MediaFileUpload(ItemPath, resumable=True)
				).execute()

			print(f"Uploaded file: {ItemName}")

for Folder in os.listdir(NewFolder):
	UploadFolder(os.path.join(NewFolder, Folder), '')

input("Finished uploading files")
